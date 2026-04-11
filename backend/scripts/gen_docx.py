#!/usr/bin/env python3
"""
gen_docx.py — Professional Word document generator for Barq Cowork.

Reads JSON from stdin:
  {
    "title":    "Document Title",
    "subtitle": "Optional subtitle",
    "author":   "Author Name",
    "date":     "April 11, 2026",
    "sections": [
      {
        "heading": "Section Heading",
        "level":   1,          // 1 = H1, 2 = H2
        "content": "Text body. Use '• item' prefix lines for bullets.",
        "table":   {           // optional
          "headers": ["Col A", "Col B"],
          "rows":    [["r1c1", "r1c2"], ...]
        }
      },
      ...
    ]
  }

Writes .docx bytes to stdout.
"""

import sys
import io
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# Theme colours
# ─────────────────────────────────────────────
DARK_BG    = RGBColor(0x0F, 0x17, 0x2A)   # near-black navy
ACCENT     = RGBColor(0x63, 0x66, 0xF1)   # indigo accent
ACCENT2    = RGBColor(0xA5, 0xB4, 0xFC)   # light indigo
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF1, 0xF5, 0xF9)
MID_GRAY   = RGBColor(0x94, 0xA3, 0xB8)
TEXT_DARK  = RGBColor(0x1E, 0x29, 0x3B)
BORDER_CLR = RGBColor(0xE2, 0xE8, 0xF0)

# ─────────────────────────────────────────────
# XML helpers
# ─────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """Set table cell background via XML shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, sides=('top','bottom','left','right'), color='E2E8F0', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_para_border_bottom(para, color='6366F1', sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def _shade_paragraph(para, hex_fill='F1F5F9'):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)


def _add_page_border(doc):
    """Add a subtle left accent bar via page borders (decorative)."""
    pass   # kept minimal — full page borders are distracting


def _run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def _run_size(run, pt: float):
    run.font.size = Pt(pt)


# ─────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────
def add_cover(doc: Document, title: str, subtitle: str, author: str, date: str):
    # Large accent top bar via paragraph shading + border
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after  = Pt(0)
    _shade_paragraph(bar, '6366F1')
    # Height trick: empty run with large font
    run = bar.add_run('\u00a0')
    run.font.size = Pt(28)
    _set_spacing(bar, before=0, after=0)

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(48)
    sp.paragraph_format.space_after  = Pt(0)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title_para.add_run(title)
    tr.font.size = Pt(36)
    tr.font.bold = True
    tr.font.color.rgb = TEXT_DARK
    _set_spacing(title_para, before=0, after=240)

    # Accent underline
    ul = doc.add_paragraph()
    _shade_paragraph(ul, '6366F1')
    ul.paragraph_format.space_before = Pt(0)
    ul.paragraph_format.space_after  = Pt(6)
    rul = ul.add_run('\u00a0')
    rul.font.size = Pt(3)
    # Make the bar ~3 inches wide via indentation trick is unreliable;
    # instead just use a short line
    _set_spacing(ul, before=0, after=120)

    # Subtitle
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sr = sub_para.add_run(subtitle)
        sr.font.size = Pt(16)
        sr.font.italic = True
        sr.font.color.rgb = MID_GRAY
        _set_spacing(sub_para, before=0, after=480)

    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Author / Date block
    if author:
        ap = doc.add_paragraph()
        ar = ap.add_run(f'Prepared by: ')
        ar.font.size = Pt(11)
        ar.font.color.rgb = MID_GRAY
        av = ap.add_run(author)
        av.font.size = Pt(11)
        av.font.bold = True
        av.font.color.rgb = TEXT_DARK
        _set_spacing(ap, before=0, after=60)

    if date:
        dp = doc.add_paragraph()
        dr = dp.add_run(f'Date: ')
        dr.font.size = Pt(11)
        dr.font.color.rgb = MID_GRAY
        dv = dp.add_run(date)
        dv.font.size = Pt(11)
        dv.font.color.rgb = TEXT_DARK
        _set_spacing(dp, before=0, after=0)

    # Page break
    doc.add_page_break()


# ─────────────────────────────────────────────
# Section heading
# ─────────────────────────────────────────────
def add_heading(doc: Document, text: str, level: int):
    if level == 1:
        # Custom styled H1 — accent left border effect via paragraph shading + bottom border
        p = doc.add_paragraph()
        _set_spacing(p, before=360, after=120)
        run = p.add_run(text.upper())
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = ACCENT
        run.font.name = 'Calibri'
        _set_para_border_bottom(p, color='6366F1', sz=6)
    else:
        # H2
        p = doc.add_paragraph()
        _set_spacing(p, before=240, after=80)
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = TEXT_DARK
        run.font.name = 'Calibri'
        _set_para_border_bottom(p, color='A5B4FC', sz=4)


# ─────────────────────────────────────────────
# Body content (text + bullets)
# ─────────────────────────────────────────────
def add_content(doc: Document, content: str):
    if not content:
        return
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith('• ') or line.startswith('- '):
            # Bullet list
            bullet_text = line[2:].strip()
            p = doc.add_paragraph()
            _set_spacing(p, before=40, after=40, line=276)
            # Manual bullet with indent
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '360')
            ind.set(qn('w:hanging'), '180')
            pPr.append(ind)
            # Bullet character
            br = p.add_run('▸  ')
            br.font.color.rgb = ACCENT
            br.font.size = Pt(10)
            tr = p.add_run(bullet_text)
            tr.font.size = Pt(10.5)
            tr.font.color.rgb = TEXT_DARK
            tr.font.name = 'Calibri'
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            _set_spacing(p, before=60, after=80, line=276)
            run = p.add_run(line)
            run.font.size = Pt(10.5)
            run.font.color.rgb = TEXT_DARK
            run.font.name = 'Calibri'
        i += 1


# ─────────────────────────────────────────────
# Table
# ─────────────────────────────────────────────
def add_table(doc: Document, headers: list, rows: list):
    if not headers:
        return

    # Convert any non-string cell to string
    def cell_str(v):
        return str(v) if v is not None else ''

    ncols = len(headers)
    nrows = len(rows)
    table = doc.add_table(rows=1 + nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        _set_cell_bg(hdr_cells[j], '6366F1')
        _set_cell_border(hdr_cells[j], color='6366F1')
        p = hdr_cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        _set_spacing(p, before=80, after=80)

    # Data rows
    for i, row in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        fill = 'FFFFFF' if i % 2 == 0 else 'F8FAFC'
        for j in range(ncols):
            val = cell_str(row[j]) if j < len(row) else ''
            _set_cell_bg(row_cells[j], fill)
            _set_cell_border(row_cells[j], color='E2E8F0')
            p = row_cells[j].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_DARK
            run.font.name = 'Calibri'
            _set_spacing(p, before=60, after=60)

    # Space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(12)


# ─────────────────────────────────────────────
# Default document styles
# ─────────────────────────────────────────────
def setup_doc_styles(doc: Document):
    # Page margins (1 inch sides, 1 inch top/bottom)
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = TEXT_DARK


def add_footer(doc: Document, title: str):
    """Add page number footer."""
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_spacing(fp, before=0, after=0)

    # Title on left, page number on right
    lr = fp.add_run(f'{title}  |  ')
    lr.font.size = Pt(8)
    lr.font.color.rgb = MID_GRAY

    # Page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run = fp.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = MID_GRAY
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ─────────────────────────────────────────────
# Main entry point
# ─────────────────────────────────────────────
def generate(data: dict) -> bytes:
    title    = data.get('title', 'Document')
    subtitle = data.get('subtitle', '')
    author   = data.get('author', '')
    date     = data.get('date', '')
    sections = data.get('sections', [])

    doc = Document()
    setup_doc_styles(doc)

    # Cover page
    add_cover(doc, title, subtitle, author, date)

    # Add footer to all sections
    add_footer(doc, title)

    # Body sections
    for sec in sections:
        heading = sec.get('heading', '')
        level   = int(sec.get('level', 1))
        content = sec.get('content', '')
        table   = sec.get('table')

        if heading:
            add_heading(doc, heading, level)
        if content:
            add_content(doc, content)
        if table and isinstance(table, dict):
            headers = table.get('headers', [])
            rows    = table.get('rows', [])
            if headers:
                add_table(doc, headers, rows)

    # Write to buffer
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == '__main__':
    raw = sys.stdin.buffer.read()
    data = json.loads(raw)
    out = generate(data)
    sys.stdout.buffer.write(out)
